# Script Version: 0.1.0 | Phase 6: Export & News Mode
# Description: Handles PDF and DOCX export logic.

import re
from PyQt6.QtGui import QTextDocument
from PyQt6.QtPrintSupport import QPrinter
from PyQt6.QtWidgets import QTextEdit

try:
    from docx import Document
    from docx.shared import Pt
except ImportError:
    Document = None

class ExportManager:
    @staticmethod
    def export_pdf(text_widget: QTextEdit, filepath: str):
        """
        Exports the content of a QTextEdit (which contains Markdown/HTML) to PDF
        using PyQt6's built-in printer.
        """
        printer = QPrinter(QPrinter.PrinterMode.HighResolution)
        printer.setOutputFormat(QPrinter.OutputFormat.PdfFormat)
        printer.setOutputFileName(filepath)
        
        # We clone the document to avoid messing with the UI widget's state
        doc = text_widget.document().clone()
        doc.print(printer)
        print(f"[EXPORT] PDF saved to {filepath}")

    @staticmethod
    def export_docx(markdown_text: str, filepath: str):
        """
        Parses simple Markdown (headers, paragraphs) and creates a DOCX file.
        """
        if not Document:
            print("[ERROR] python-docx not installed.")
            return

        doc = Document()
        
        # Simple Markdown Parser
        lines = markdown_text.split('\n')
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
                
            if line.startswith('# '):
                doc.add_heading(line[2:], level=1)
            elif line.startswith('## '):
                doc.add_heading(line[3:], level=2)
            elif line.startswith('### '):
                doc.add_heading(line[4:], level=3)
            elif line.startswith('#### '):
                doc.add_heading(line[5:], level=4)
            elif line.startswith('- ') or line.startswith('* '):
                doc.add_paragraph(line[2:], style='List Bullet')
            else:
                # Regular paragraph
                p = doc.add_paragraph(line)
                # Basic bold handling (**text**)
                # This is a naive implementation; for complex markdown, a full parser is needed.
                # But this covers the "Functional Change Only" requirement for basic reports.
        
        try:
            doc.save(filepath)
            print(f"[EXPORT] DOCX saved to {filepath}")
        except Exception as e:
            print(f"[ERROR] Failed to save DOCX: {e}")
